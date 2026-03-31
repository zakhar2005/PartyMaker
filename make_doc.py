from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Поля страницы
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)

# Стили
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(14)

def heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    p.paragraph_format.space_after = Pt(0)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    return p

def code_line(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(0)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    for r, row in enumerate(rows):
        cells = table.rows[r + 1].cells
        for i, val in enumerate(row):
            cells[i].text = val
            for run in cells[i].paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    doc.add_paragraph()

# ── ЗАГОЛОВОК ДОКУМЕНТА ──────────────────────────────────────────────────────

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('PartyMaker — документация')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(16)

# ── РАЗДЕЛЫ ───────────────────────────────────────────────────────────────────

heading('Общее описание проекта')
body(
    'PartyMaker — это веб-приложение для совместного выбора фильма в группе. '
    'Сервис показывает участникам фильмы из каталога TMDB, собирает лайки и дизлайки каждого, '
    'после первого раунда вычисляет жанровые предпочтения группы и предлагает следующие фильмы '
    'с учётом этих предпочтений. Голосование продолжается, пока все участники не поставят лайк '
    'одному и тому же фильму.', indent=True
)

heading('Назначение системы')
body('Система предназначена для:')
bullet('демонстрации реализации клиентского SPA без серверной части и сторонних фреймворков;')
bullet('интеграции с внешними REST API (TMDB, QR-сервер);')
bullet('реализации алгоритма жанровой фильтрации и поиска группового консенсуса;')
bullet('хранения пользовательских данных в localStorage браузера.')

heading('Основные функции')
bullet('создание игровой комнаты с уникальным кодом-идентификатором;')
bullet('добавление участников и отслеживание их голосов;')
bullet('голосование лайк/дизлайк по 3 фильмам за раунд с отображением прогресса;')
bullet('расчёт топ-3 жанров по итогам первого раунда и жанровая фильтрация последующих;')
bullet('определение победителя при единогласном лайке всех участников;')
bullet('генерация QR-кода со ссылкой на страницу победившего фильма на TMDB;')
bullet('сохранение победителей в локальный архив и режим «Вечер воспоминаний» для повторного голосования.')

heading('Используемые технологии')
bullet('язык программирования: JavaScript (ES6+);')
bullet('интерфейс: HTML5, CSS3 (без фреймворков и сборки);')
bullet('внешний каталог фильмов: TMDB API;')
bullet('генерация QR-кодов: api.qrserver.com;')
bullet('локальное хранилище данных: localStorage;')
bullet('веб-сервер для локального запуска: Python http.server.')

heading('Процедура запуска')
body('Для запуска приложения необходимо выполнить команду в директории проекта:')
code_line('python3 -m http.server 4173 --bind 0.0.0.0')
body('После этого открыть в браузере адрес: http://localhost:4173', indent=True)
body(
    'Для доступа с других устройств в локальной сети следует использовать '
    'IP-адрес машины вместо localhost.', indent=True
)
body('Конфигурация задаётся в файле app.js двумя константами:')
code_line("const TMDB_API_KEY = '';       // ключ TMDB для загрузки реального каталога")
code_line('const FIREBASE_CONFIG = null;  // зарезервировано для многоустройственного режима')
body(
    'Без TMDB_API_KEY приложение работает с встроенным списком из 20 фильмов.', indent=True
)

heading('Краткое описание алгоритма')
body(
    'Раунд 1. Показываются первые 3 фильма из пула. После голосования всех участников '
    'для каждого жанра вычисляется вес:', indent=True
)
code_line('вес(жанр) = сумма (лайки − дизлайки) по всем фильмам с этим жанром')
body('Из результата берутся топ-3 жанра группы.', indent=True)
body(
    'Раунды 2 и далее. Для каждого непоказанного фильма вычисляется оценка:', indent=True
)
code_line('score = сумма (3 − позиция_жанра_в_топе) для каждого совпавшего жанра')
body(
    'Фильмы с нулевым score исключаются. Показываются 3 фильма с наибольшим score; '
    'при равенстве — с наибольшим рейтингом.', indent=True
)
body(
    'Завершение. Если хотя бы один фильм получил лайк от всех участников — он объявляется '
    'победителем и генерируется QR-код. Если пул исчерпан без консенсуса — выводится '
    'соответствующее сообщение.', indent=True
)

heading('Описание функций')
body('Внешние API-запросы:')
add_table(
    ['Метод', 'URL', 'Назначение'],
    [
        ['GET', 'api.themoviedb.org/3/genre/movie/list', 'Список жанров с ID для маппинга'],
        ['GET', 'api.themoviedb.org/3/movie/top_rated',  'Страница топ-рейтинговых фильмов'],
        ['GET', 'api.qrserver.com/v1/create-qr-code/',  'Генерация PNG QR-кода по переданному URL'],
    ]
)
body('Основные функции приложения:')
add_table(
    ['Функция', 'Назначение'],
    [
        ['createRoom()',         'Создаёт комнату, генерирует код, сбрасывает состояние'],
        ['joinGuest()',          'Добавляет участника в комнату вручную'],
        ['startVoting()',        'Запускает первый раунд; в режиме воспоминаний подставляет архивный пул'],
        ['processRound()',       'Ищет консенсус, вычисляет жанры, готовит следующий батч'],
        ['vote()',               'Записывает голос участника в state и перерисовывает доску'],
        ['getNextBatch()',       'Отбирает и ранжирует следующие 3 фильма по жанровым предпочтениям'],
        ['computeTopGenres…()', 'Вычисляет топ-3 жанра по итогам первого раунда'],
        ['findUnanimousMovie()', 'Проверяет наличие фильма с единогласным лайком'],
        ['showFinal()',          'Фиксирует победителя, сохраняет в архив, отображает результат'],
        ['renderBoard()',        'Перерисовывает карточки фильмов с кнопками голосования'],
        ['loadMoviesFromTMDB()', 'Загружает каталог с TMDB, маппит жанры по ID'],
        ['saveWinner() / loadWinners() / clearWinners()', 'Управление архивом победителей в localStorage'],
    ]
)

heading('Описание внесённых изменений в исходный вариант бизнес-логики')

body(
    '1. Загрузка фильмов из TMDB вместо захардкоженного списка. '
    'Исходный вариант содержал массив из 20 фильмов, заданных вручную. В текущей версии '
    'при наличии API-ключа приложение загружает реальный каталог (~57 фильмов) с русскими '
    'названиями и жанрами через TMDB API. При отсутствии ключа или ошибке сети используется '
    'встроенный список.', indent=True
)
body(
    '2. Режим «Вечер воспоминаний». Добавлен флаг state.memoryMode. '
    'При его активации пул фильмов формируется из архива победителей предыдущих сессий, '
    'хранящегося в localStorage. Победитель каждой сессии автоматически добавляется в архив.', indent=True
)
body(
    '3. Контроль завершённости голосования. В исходном варианте кнопка «Обработать раунд» '
    'была доступна в любой момент. В текущей версии кнопка заблокирована до тех пор, пока '
    'все участники не проголосуют по всем фильмам. В тексте кнопки отображается счётчик '
    'заполненных голосов.', indent=True
)
body(
    '4. QR-код итогового фильма. В исходном варианте QR кодировал JSON-объект, '
    'который при сканировании телефоном отображался как нечитаемая строка. В текущей версии '
    'QR кодирует прямую ссылку на страницу фильма на TMDB.', indent=True
)
body(
    '5. UX-улучшения. Добавлена поддержка клавиши Enter в полях ввода. Кнопки получили '
    'состояния hover, active и disabled с визуальной обратной связью. Поля ввода '
    'выделяются при фокусе.', indent=True
)

heading('Возможные направления развития')
body(
    'В качестве дальнейшего развития проекта можно рассмотреть добавление '
    'многоустройственного режима: синхронизацию состояния голосования между несколькими '
    'устройствами в реальном времени через Firebase Realtime Database. В таком сценарии '
    'хост управляет комнатой со своего устройства, а участники голосуют с телефонов, '
    'подключившись по коду комнаты.', indent=True
)

heading('Вывод')
body(
    'Проект PartyMaker демонстрирует практическую реализацию клиентского веб-приложения '
    'с нетривиальной бизнес-логикой: итеративный алгоритм жанровой фильтрации, интеграция '
    'с внешними API и управление состоянием без фреймворков. Приложение не требует серверной '
    'части и разворачивается как статический сайт на любом хостинге.', indent=True
)

doc.save('Документация_PartyMaker.docx')
print('Готово: Документация_PartyMaker.docx')
